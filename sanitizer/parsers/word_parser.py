"""
Word 文档解析器
支持 .docx 格式
"""

import logging
from pathlib import Path

from sanitizer.engine import SanitizeEngine
from sanitizer.parsers.base import BaseParser

logger = logging.getLogger(__name__)


class WordParser(BaseParser):
    """
    Word 文档解析器

    使用 python-docx 遍历文档中的：
    - paragraphs（段落）中的每个 run
    - tables（表格）中的每个 cell
    保留原始格式（字体、颜色、大小等）
    """

    def can_handle(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() == ".docx"

    def sanitize(self, input_path: str, output_path: str, engine: SanitizeEngine) -> dict:
        from docx import Document

        engine.set_current_file(input_path)
        stats_before = dict(engine.stats)

        # 数值缩放用的 seed（按文件名）
        scale_seed = Path(input_path).name

        doc = Document(input_path)

        # 处理所有段落
        for paragraph in doc.paragraphs:
            self._process_paragraph(paragraph, engine, scale_seed)

        # 处理所有表格
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._process_paragraph(paragraph, engine, scale_seed)

        # 处理页眉和页脚
        for section in doc.sections:
            # 页眉
            if section.header and section.header.paragraphs:
                for paragraph in section.header.paragraphs:
                    self._process_paragraph(paragraph, engine, scale_seed)
            # 页脚
            if section.footer and section.footer.paragraphs:
                for paragraph in section.footer.paragraphs:
                    self._process_paragraph(paragraph, engine, scale_seed)

        # 保存
        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        doc.save(output_path)

        return {
            "entities_found": engine.stats["entities_found"] - stats_before["entities_found"],
            "entities_replaced": engine.stats["entities_replaced"] - stats_before["entities_replaced"],
        }

    @staticmethod
    def _process_paragraph(paragraph, engine: SanitizeEngine, scale_seed: str = ""):
        """
        处理单个段落

        策略：
        1. 先把段落内所有 run 的文本拼合为完整字符串
        2. 对完整字符串调用引擎脱敏（实体替换 + 数值缩放）
           避免实体被 run 边界截断而漏识别
        3. 如果文本发生了变化：
           - 将脱敏后的完整文本写入第一个非空 run
           - 清空其余 run 的文本
           （保留第一个 run 的格式；其余 run 的格式丢失，但内容正确）
        4. 如果文本没有变化，所有 run 保持不变（格式完全保留）
        """
        if not paragraph.runs:
            return

        # 拼合完整段落文本
        full_text = "".join(run.text for run in paragraph.runs)
        if not full_text.strip():
            return

        # 第一步：实体脱敏（公司名、邮箱、电话等）
        sanitized = engine.process_text(full_text)

        # 第二步：数值缩放（金额、数量等）
        if scale_seed:
            sanitized = engine.scale_numbers_in_text(sanitized, scale_seed)

        # 没有变化则不做任何修改（格式完全保留）
        if sanitized == full_text:
            return

        # 文本发生了变化：写入第一个 run，清空其余
        first_written = False
        for run in paragraph.runs:
            if not first_written:
                run.text = sanitized
                first_written = True
            else:
                run.text = ""
